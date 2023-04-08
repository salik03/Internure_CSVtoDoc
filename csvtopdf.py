import pandas as pd
import docx

data = pd.read_csv('resume_data.csv')

for i in range(len(data)):
    document = docx.Document()
    heading = document.add_paragraph(style='Heading1')
    heading.add_run(data['Full Name'][i])
    heading.add_run('\n')
    heading.add_run(data['Email'][i])
    heading.add_run(' | ')
    heading.add_run(str(data['Phone Number'][i]))
    heading.add_run('\n')
    heading.add_run(data['LinkedIn URL'][i])
    heading.add_run(' | ')
    heading.add_run(data['Github URL'][i])
    document.add_paragraph('Skills', style='Heading2')
    skills_list = document.add_paragraph()
    for skill in ['Languages', 'Frameworks', 'Tools']:
        skills_list.add_run(f'{skill}: ')
        skills_list.add_run(data[skill][i]).bold = True
        skills_list.add_run('\n')
    document.add_paragraph('Education', style='Heading2')
    p = document.add_paragraph(style='List Bullet')
    p.add_run(data['Education'][i]).bold = True
    p.add_run(f' ({data["Degree"][i]} in {data["Major"][i]})')
    p.add_run(f', {data["Graduation Date"][i]}')
    document.add_paragraph('Experience', style='Heading2')
    p = document.add_paragraph(style='List Bullet')
    p.add_run(data['Experience'][i]).bold = True
    p.add_run(f', {data["Company"][i]}')
    p.add_run(f' ({data["Start Date"][i]} - {data["End Date"][i]})')
    p = document.add_paragraph(style='Body Text')
    p.add_run(data['Description'][i])
    document.add_paragraph('Projects', style='Heading2')
    p = document.add_paragraph(style='List Bullet')
    p.add_run(data['Projects'][i]).bold = True
    p = document.add_paragraph(style='Body Text')
    p.add_run(data['Project Description'][i])
    document.save(f'{data["Full Name"][i]}.docx')
